
import sys
import os
import argparse
import markdown
from docx import Document
from docx.shared import Pt, Mm, RGBColor
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from bs4 import BeautifulSoup, Tag

# Import from md_to_docx package
from md_to_docx.config import load_config
from md_to_docx.utils import fix_markdown_tables
from md_to_docx.formatters import (
    set_cell_background,
    set_cell_margins,
    set_table_border,
    set_cell_border,
    apply_heading_format
)

# Load configuration
config = load_config()

# Parse command-line arguments
parser = argparse.ArgumentParser(description='Convert Markdown to Word document')
parser.add_argument('input_md', help='Input Markdown file')
parser.add_argument('output_docx', nargs='?', help='Output Word document file (optional, defaults to input name with .docx extension)')
parser.add_argument('--no-show', action='store_true', help='Do not open the file after creation')
args = parser.parse_args()

input_md = args.input_md

# If output file not specified, use input filename with .docx extension
if args.output_docx:
    output_docx = args.output_docx
else:
    base_name = os.path.splitext(input_md)[0]
    output_docx = f"{base_name}.docx"

# Check if input file exists
if not os.path.exists(input_md):
    print(f"Error: Input file '{input_md}' not found.")
    sys.exit(1)

# Check if output file exists
if os.path.exists(output_docx):
    response = input(f"{output_docx} already exists. Delete and overwrite? (y/n): ")
    if response.lower() != 'y':
        print("Operation cancelled.")
        sys.exit(0)
    try:
        os.remove(output_docx)
    except PermissionError:
        print(f"Error: Cannot delete {output_docx}. The file may be open in another program.")
        print("Please close the file and try again.")
        sys.exit(1)

# Load Markdown content
with open(input_md, "r", encoding="utf-8") as f:
    md_content = f.read()

# Fix markdown tables with blank lines between rows
md_content = fix_markdown_tables(md_content)

# Convert Markdown to HTML
html = markdown.markdown(md_content, extensions=["fenced_code", "tables"])

# Parse HTML
soup = BeautifulSoup(html, "html.parser")

# Create Word document
doc = Document()

# Set default font to Calibri
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'

# Process all elements (use find_all to get all elements at any level)
for element in soup.find_all(recursive=False):
    if isinstance(element, Tag):
        if element.name == "p":
            text = element.get_text()
            paragraph_config = config.get('paragraph', {})
            
            paragraph = doc.add_paragraph(text)
            if paragraph.runs:
                run = paragraph.runs[0]
            else:
                run = paragraph.add_run()
            
            # Apply font settings
            run.font.name = paragraph_config.get('font_name', 'Calibri')
            run.font.size = Pt(paragraph_config.get('font_size', 11))
            run.font.bold = paragraph_config.get('bold', False)
            run.font.italic = paragraph_config.get('italic', False)
            
            # Apply color if specified
            color = paragraph_config.get('color')
            if color:
                if len(color) == 6:
                    r = int(color[0:2], 16)
                    g = int(color[2:4], 16)
                    b = int(color[4:6], 16)
                    run.font.color.rgb = RGBColor(r, g, b)
            
            # Apply spacing and line spacing
            paragraph.paragraph_format.space_before = Pt(paragraph_config.get('space_before', 0))
            paragraph.paragraph_format.space_after = Pt(paragraph_config.get('space_after', 8))
            paragraph.paragraph_format.line_spacing = paragraph_config.get('line_spacing', 1.15)
        elif element.name in ["h1", "h2", "h3", "h4", "h5", "h6"]:
            # Add headers with configured formatting
            heading_name = element.name
            text = element.get_text()
            
            # Check if we have config for this heading level
            if heading_name in config.get('headings', {}):
                # Add paragraph and apply custom formatting
                paragraph = doc.add_paragraph(text)
                apply_heading_format(paragraph, config['headings'][heading_name])
            else:
                # Fall back to default heading style
                level = int(element.name[1])
                doc.add_heading(text, level=level)
        elif element.name == "pre":
            code_text = element.get_text().strip()
            # Get codeblock config
            codeblock_config = config.get('codeblock', {})
            
            # Create a table with one cell for the code block
            table = doc.add_table(rows=1, cols=1)
            cell = table.cell(0, 0)
            
            # Add code text with monospace font
            paragraph = cell.paragraphs[0]
            run = paragraph.add_run(code_text)
            run.font.name = codeblock_config.get('font_name', 'Courier New')
            run.font.size = Pt(codeblock_config.get('font_size', 10))
            
            # Remove paragraph spacing
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.line_spacing = codeblock_config.get('line_spacing', 1.0)
            
            # Set cell margins
            set_cell_margins(cell, Mm(codeblock_config.get('cell_margin', 3)))
            
            # Set cell background
            set_cell_background(cell, codeblock_config.get('background_color', 'F2F2F2'))
            
            # Set 1pt black border
            set_table_border(table)
            
            # Add empty paragraph after table for spacing
            doc.add_paragraph()
        elif element.name == "ul":
            # Handle unordered lists
            for li in element.find_all("li", recursive=False):
                doc.add_paragraph(li.get_text(), style='List Bullet')
        elif element.name == "ol":
            # Handle ordered lists
            for li in element.find_all("li", recursive=False):
                doc.add_paragraph(li.get_text(), style='List Number')
        elif element.name == "table":
            # Handle tables
            table_config = config.get('table', {})
            header_config = table_config.get('header', {})
            data_config = table_config.get('data', {})
            
            rows = element.find_all("tr")
            if not rows:
                continue
                
            # Count maximum columns
            max_cols = 0
            for row in rows:
                cols = len(row.find_all(["th", "td"]))
                max_cols = max(max_cols, cols)
            
            if max_cols == 0:
                continue
                
            # Create Word table
            table = doc.add_table(rows=len(rows), cols=max_cols)
            
            # Process each row
            for row_idx, row in enumerate(rows):
                cells = row.find_all(["th", "td"])
                for col_idx, cell in enumerate(cells):
                    if col_idx < max_cols:
                        word_cell = table.cell(row_idx, col_idx)
                        
                        # Clear paragraph and set up formatting
                        paragraph = word_cell.paragraphs[0]
                        paragraph.clear()
                        
                        # Set 0pt spacing after paragraph
                        paragraph.paragraph_format.space_after = Pt(0)
                        
                        # Handle header row (first row or th elements)
                        if row_idx == 0 or cell.name == "th":
                            cell_text = cell.get_text().strip()
                            # Apply uppercase if configured
                            if header_config.get('uppercase', True):
                                cell_text = cell_text.upper()
                            
                            # Header formatting from config
                            run = paragraph.add_run(cell_text)
                            run.font.bold = header_config.get('bold', True)
                            run.font.size = Pt(header_config.get('font_size', 10))
                            run.font.name = header_config.get('font_name', 'Calibri')
                            
                            # Set vertical alignment for header cells
                            alignment = header_config.get('vertical_alignment', 'bottom')
                            if alignment.lower() == 'bottom':
                                word_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.BOTTOM
                            elif alignment.lower() == 'center':
                                word_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                            elif alignment.lower() == 'top':
                                word_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
                            
                            # Header row borders: no borders except bottom
                            set_cell_border(word_cell, {
                                'top': False,
                                'left': False,
                                'bottom': True,
                                'right': False
                            })
                        else:
                            # Data row formatting - handle styled spans
                            # Check if cell contains span elements with color styles
                            span_elements = cell.find_all('span')
                            if span_elements:
                                # Process all descendants but skip text nodes inside span elements
                                # to avoid duplication (the span's text will be captured when processing the span)
                                for elem in cell.descendants:
                                    if isinstance(elem, str):
                                        # Only process text nodes that are NOT direct children of span elements
                                        if elem.strip() and (not hasattr(elem, 'parent') or elem.parent.name != 'span'):
                                            run = paragraph.add_run(elem)
                                            run.font.name = data_config.get('font_name', 'Calibri')
                                    elif elem.name == 'span':
                                        # Check for color style
                                        style = elem.get('style', '')
                                        text = elem.get_text()
                                        run = paragraph.add_run(text)
                                        run.font.name = data_config.get('font_name', 'Calibri')
                                        
                                        # Parse color from style attribute
                                        if 'color:' in style:
                                            # Extract color value (e.g., "color:red;" or "color:#FF0000;")
                                            color_part = [s for s in style.split(';') if 'color:' in s]
                                            if color_part:
                                                color_value = color_part[0].split(':')[1].strip()
                                                
                                                # Map common color names to RGB
                                                color_map = {
                                                    'red': (255, 0, 0),
                                                    'green': (0, 128, 0),
                                                    'blue': (0, 0, 255),
                                                    'orange': (255, 165, 0),
                                                    'yellow': (255, 255, 0),
                                                    'purple': (128, 0, 128),
                                                    'black': (0, 0, 0),
                                                    'white': (255, 255, 255),
                                                    'gray': (128, 128, 128),
                                                    'grey': (128, 128, 128),
                                                }
                                                
                                                if color_value.lower() in color_map:
                                                    rgb = color_map[color_value.lower()]
                                                    run.font.color.rgb = RGBColor(*rgb)
                                                elif color_value.startswith('#'):
                                                    # Handle hex color codes
                                                    hex_color = color_value.lstrip('#')
                                                    if len(hex_color) == 6:
                                                        r = int(hex_color[0:2], 16)
                                                        g = int(hex_color[2:4], 16)
                                                        b = int(hex_color[4:6], 16)
                                                        run.font.color.rgb = RGBColor(r, g, b)
                            else:
                                # No styled spans, just add plain text
                                cell_text = cell.get_text().strip()
                                run = paragraph.add_run(cell_text)
                                run.font.name = data_config.get('font_name', 'Calibri')
                            
                            # Data row borders: top, left, right, bottom for data rows
                            borders = {
                                'top': row_idx == 1,  # Only first data row gets top border
                                'left': True,
                                'bottom': True,
                                'right': True
                            }
                            set_cell_border(word_cell, borders)
            
            # Add empty paragraph after table for spacing
            doc.add_paragraph()

try:
    doc.save(output_docx)
    print(f"Successfully created {output_docx}")
    
    # Open the file unless --no-show is specified
    if not args.no_show:
        os.startfile(output_docx)
except PermissionError:
    print(f"Error: Cannot save {output_docx}. The file may be open in another program.")
    print("Please close the file and try again.")
    sys.exit(1)
except Exception as e:
    print(f"Error saving file: {e}")
    sys.exit(1)
